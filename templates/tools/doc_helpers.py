"""
Shared helpers for generating branded Nakola Expert Systems Word documents.
Keeps every generated document visually consistent: same letterhead header,
footer, heading styles, table styling, and placeholder highlighting.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand constants ──────────────────────────────────────────────
BRAND_BLUE = RGBColor(0x25, 0x63, 0xEB)
BRAND_PURPLE = RGBColor(0x7C, 0x3A, 0xED)
BRAND_DARK = RGBColor(0x0A, 0x0F, 0x1A)
BODY_TEXT = RGBColor(0x1F, 0x29, 0x37)
GRAY = RGBColor(0x64, 0x74, 0x8B)
LIGHT_BLUE_RULE = "DBEAFE"
PLACEHOLDER_COLOR = RGBColor(0xB4, 0x2C, 0x0C)
LEGAL_NOTICE_BG = "FEF3C7"

FONT_NAME = "Calibri"
ASSETS_DIR = os.path.join(os.path.dirname(__file__), "assets")
ICON_PATH = os.path.join(ASSETS_DIR, "nes-icon.png")

COMPANY = {
    "name": "Nakola Expert Systems",
    "tagline": "Build Smarter. Scale Faster. Innovate Globally.",
    "email": "hello@nakolaexpertsystems.com",
    "phone": "+254 715 674 828",
    "website": "www.nakolaexpertsystems.com",
    "location": "Nairobi, Kenya",
    "founder": "Lucky Nakola",
    "founder_title": "Founder",
}

PLACEHOLDER_LEGEND = (
    "Fields shown in orange italics are official records not yet published "
    "(registration numbers, banking details, etc.) — fill these in before "
    "sending this document, and delete this note."
)


def _set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_bottom_border(paragraph, color="2563EB", size=8, space=4):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def new_document():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(10.5)
    style.font.color.rgb = BODY_TEXT
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)

    section = doc.sections[0]
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    add_letterhead_header(doc)
    add_footer(doc)
    return doc


def add_letterhead_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    table = header.add_table(rows=1, cols=2, width=Inches(6.9))
    table.autofit = True
    left, right = table.rows[0].cells
    left.width = Inches(1.0)
    right.width = Inches(5.9)

    p = left.paragraphs[0]
    run = p.add_run()
    run.add_picture(ICON_PATH, height=Inches(0.5))

    p2 = right.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p2.add_run(COMPANY["name"])
    r1.bold = True
    r1.font.size = Pt(15)
    r1.font.color.rgb = BRAND_DARK

    p3 = right.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p3.add_run(COMPANY["tagline"])
    r2.italic = True
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = BRAND_PURPLE

    rule_p = header.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(4)
    _add_bottom_border(rule_p, color="2563EB", size=10, space=2)


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        f"{COMPANY['name']}  •  {COMPANY['email']}  •  {COMPANY['phone']}  •  {COMPANY['website']}   |   Page "
    )
    run.font.size = Pt(8)
    run.font.color.rgb = GRAY
    _add_page_number_field(p)
    for r in p.runs[1:]:
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY


def doc_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = BRAND_DARK
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(subtitle)
        r2.font.size = Pt(11)
        r2.font.color.rgb = GRAY
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_after = Pt(12)
    _add_bottom_border(rule_p, color="E2E8F0", size=6, space=2)


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BRAND_BLUE
    _add_bottom_border(p, color=LIGHT_BLUE_RULE, size=6, space=3)
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = BRAND_DARK
    return p


def body(doc, text, *, bold=False, italic=False, size=10.5, color=None, space_after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else BODY_TEXT
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    return p


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    return p


def field(doc, label, value, is_placeholder=False):
    """A single 'Label: value' line. Placeholder values are styled to stand out."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = BODY_TEXT
    r2 = p.add_run(value)
    r2.font.size = Pt(10.5)
    if is_placeholder:
        r2.italic = True
        r2.font.color.rgb = PLACEHOLDER_COLOR
    else:
        r2.font.color.rgb = BODY_TEXT
    return p


def placeholder_legend(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run(f"⚠ {PLACEHOLDER_LEGEND}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = PLACEHOLDER_COLOR


def legal_notice(doc, text="This is a business template and must be reviewed by a qualified "
                            "legal professional before official use."):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_background(cell, LEGAL_NOTICE_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"⚖  Legal Notice: {text}")
    run.bold = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def make_table(doc, headers, rows, col_widths=None, align_right_cols=None):
    align_right_cols = align_right_cols or set()
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        _set_cell_background(hdr_cells[i], "2563EB")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        if r_i % 2 == 1:
            for c in cells:
                _set_cell_background(c, "F1F5F9")
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in align_right_cols else WD_ALIGN_PARAGRAPH.LEFT
            is_ph = isinstance(val, str) and val.startswith('[') and val.endswith(']')
            run = p.add_run(val)
            run.font.size = Pt(9.5)
            if is_ph:
                run.italic = True
                run.font.color.rgb = PLACEHOLDER_COLOR
            else:
                run.font.color.rgb = BODY_TEXT

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(8)
    return table


def signature_block(doc, left_label, left_name, right_label=None, right_name=None):
    table = doc.add_table(rows=4, cols=2 if right_label else 1)
    table.autofit = True
    labels = [
        ("Name / Title:", left_name),
        ("Signature:", "________________________"),
        ("Date:", "________________________"),
        ("", ""),
    ]
    left_header = table.rows[0].cells[0].paragraphs[0].add_run(left_label)
    left_header.bold = True
    left_header.font.color.rgb = BRAND_BLUE
    left_header.font.size = Pt(10.5)

    if right_label:
        right_header = table.rows[0].cells[1].paragraphs[0].add_run(right_label)
        right_header.bold = True
        right_header.font.color.rgb = BRAND_BLUE
        right_header.font.size = Pt(10.5)

    row_data = [
        (f"Name / Title: {left_name}", f"Name / Title: {right_name}" if right_name else None),
        ("Signature: ________________________", "Signature: ________________________" if right_label else None),
        ("Date: ________________________", "Date: ________________________" if right_label else None),
    ]
    for r_i, (left_text, right_text) in enumerate(row_data, start=1):
        p = table.rows[r_i].cells[0].paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run(left_text)
        run.font.size = Pt(10)
        is_ph = '[' in left_text
        if is_ph:
            run.italic = True
            run.font.color.rgb = PLACEHOLDER_COLOR
        if right_label and right_text:
            p2 = table.rows[r_i].cells[1].paragraphs[0]
            p2.paragraph_format.space_before = Pt(6)
            run2 = p2.add_run(right_text)
            run2.font.size = Pt(10)
            if '[' in right_text:
                run2.italic = True
                run2.font.color.rgb = PLACEHOLDER_COLOR
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def save(doc, relative_path):
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
    full_path = os.path.join(repo_root, relative_path)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    doc.save(full_path)
    print(f"Generated: {relative_path}")
