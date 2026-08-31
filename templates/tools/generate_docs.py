"""
Generates professional, branded Word (.docx) versions of the 9 core
business + legal templates from templates/business-documents/ and
templates/legal/. Run this whenever brand details or template content
change, to regenerate the .docx files.

Usage:
    pip install python-docx pillow
    python templates/tools/generate_docs.py
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from doc_helpers import (
    COMPANY, new_document, doc_title, h2, h3, body, bullet, numbered,
    field, placeholder_legend, legal_notice, make_table, signature_block, save,
)

PH = lambda s: f"[{s}]"  # marks a value as an official-record placeholder


# ─────────────────────────────────────────────────────────────────
# 1. COMPANY PROFILE
# ─────────────────────────────────────────────────────────────────
def company_profile():
    doc = new_document()
    doc_title(doc, "Company Profile", "Nakola Expert Systems")

    h2(doc, "Company Information")
    field(doc, "Company Name", COMPANY["name"])
    field(doc, "Tagline", "Build Smarter. Scale Faster. Innovate Globally.")
    field(doc, "Positioning", "African-grounded technical expertise delivering global-standard "
                              "digital solutions that drive measurable business outcomes.")
    field(doc, "Founded", "2021")
    field(doc, "Headquarters", COMPANY["location"])
    field(doc, "Service Reach", "Clients across Africa and global markets")
    field(doc, "Company Registration Number", PH("Company Registration Number"), True)
    field(doc, "KRA PIN", PH("KRA PIN"), True)

    h2(doc, "About Us")
    body(doc, "Nakola Expert Systems is a software and technology partner for startups, SMEs, "
              "enterprises, and international organizations. We combine African market "
              "understanding with global engineering standards to deliver reliable, scalable "
              "digital products.")

    h2(doc, "Mission")
    body(doc, "To empower businesses across Africa and beyond with technology solutions that "
              "are expertly built, outcome-focused, and accessible.")

    h2(doc, "Vision")
    body(doc, "To become Africa's most trusted and recommended technology partner.")

    h2(doc, "Core Values")
    for v in ["Outcomes Over Output", "Radical Transparency", "Senior-Led, Always",
              "African-Global Mindset", "Continuous Learning", "Long-Term Partnership"]:
        bullet(doc, v)

    h2(doc, "Core Services")
    for s in ["Custom Software Development", "Technology Consulting", "Digital Transformation",
              "Cloud and DevOps", "Quality Assurance and Testing", "Managed Support and Maintenance"]:
        bullet(doc, s)

    h2(doc, "Proof of Delivery")
    bullet(doc, "50+ projects delivered")
    bullet(doc, "8+ countries served")
    bullet(doc, "98% client satisfaction")
    bullet(doc, "Multi-industry experience: e-commerce, healthtech, fintech, enterprise, "
                "education, and insurance")

    h2(doc, "Contact")
    field(doc, "Email", COMPANY["email"])
    field(doc, "Phone", COMPANY["phone"])
    field(doc, "Location", COMPANY["location"])
    field(doc, "Website", COMPANY["website"])

    h2(doc, "Authorized Signatory")
    field(doc, "Name", COMPANY["founder"])
    field(doc, "Title", COMPANY["founder_title"])
    field(doc, "Date", PH("Date"), True)

    save(doc, "templates/business-documents/COMPANY_PROFILE.docx")


# ─────────────────────────────────────────────────────────────────
# 2. LETTERHEAD
# ─────────────────────────────────────────────────────────────────
def letterhead():
    doc = new_document()
    doc_title(doc, "Official Correspondence")

    field(doc, "Registered Address", PH("Registered Address"), True)
    field(doc, "Nairobi Office", COMPANY["location"])
    field(doc, "Email", COMPANY["email"])
    field(doc, "Phone", COMPANY["phone"])
    field(doc, "Website", COMPANY["website"])
    field(doc, "Company Registration Number", PH("Company Registration Number"), True)
    field(doc, "KRA PIN", PH("KRA PIN"), True)

    h2(doc, "Document Metadata")
    field(doc, "Date", PH("Date"), True)
    field(doc, "Reference Number", PH("Reference Number"), True)
    field(doc, "Recipient Name", PH("Recipient Name"), True)
    field(doc, "Recipient Company", PH("Recipient Company"), True)
    field(doc, "Recipient Address", PH("Recipient Address"), True)

    h2(doc, "Subject")
    body(doc, PH("Subject Line"), italic=True, color=None)

    h2(doc, "Letter Body")
    body(doc, f"Dear {PH('Recipient Name')},", italic=False)
    body(doc, "")
    body(doc, PH("Write your message content here."), italic=True)
    body(doc, "")
    body(doc, "Sincerely,")
    body(doc, PH("Sender Name"), italic=True, color=None)
    body(doc, PH("Sender Title"), italic=True, color=None)
    body(doc, COMPANY["name"])

    placeholder_legend(doc)
    body(doc, f"This document is issued by {COMPANY['name']}. For verification, "
              f"contact {COMPANY['email']}.", size=9, color=None, italic=True)

    save(doc, "templates/business-documents/LETTERHEAD.docx")


# ─────────────────────────────────────────────────────────────────
# 3. QUOTATION
# ─────────────────────────────────────────────────────────────────
def quotation():
    doc = new_document()
    doc_title(doc, "Quotation")
    placeholder_legend(doc)

    h2(doc, "Quotation Details")
    field(doc, "Quotation Number", PH("Quotation Number"), True)
    field(doc, "Date Issued", PH("Date"), True)
    field(doc, "Valid Until", PH("Valid Until Date"), True)
    field(doc, "Currency", PH("KES / USD / Other"), True)

    h2(doc, "Company Details")
    field(doc, "Company", COMPANY["name"])
    field(doc, "Email", COMPANY["email"])
    field(doc, "Phone", COMPANY["phone"])
    field(doc, "Address", PH("Registered Address"), True)
    field(doc, "Company Registration Number", PH("Company Registration Number"), True)
    field(doc, "KRA PIN", PH("KRA PIN"), True)

    h2(doc, "Client Details")
    field(doc, "Client Name", PH("Client Name"), True)
    field(doc, "Company", PH("Client Company"), True)
    field(doc, "Email", PH("Client Email"), True)
    field(doc, "Phone", PH("Client Phone"), True)
    field(doc, "Address", PH("Client Address"), True)

    h2(doc, "Project Summary")
    field(doc, "Project Name", PH("Project Name"), True)
    field(doc, "Service Category", PH("Service Needed"), True)
    field(doc, "Scope Summary", PH("Short Scope Summary"), True)
    field(doc, "Delivery Timeline", PH("Estimated Timeline"), True)

    h2(doc, "Cost Breakdown")
    make_table(
        doc,
        ["Item", "Description", "Qty", "Unit Cost", "Line Total"],
        [
            ["1", "[Workstream 1]", "[Qty]", "[Amount]", "[Amount]"],
            ["2", "[Workstream 2]", "[Qty]", "[Amount]", "[Amount]"],
            ["3", "[Workstream 3]", "[Qty]", "[Amount]", "[Amount]"],
        ],
        align_right_cols={2, 3, 4},
    )
    field(doc, "Subtotal", PH("Amount"), True)
    field(doc, "VAT / Tax", f"({PH('Rate')}%) {PH('Amount')}", True)
    field(doc, "Discount (if any)", PH("Amount"), True)
    p = doc.add_paragraph()
    r1 = p.add_run("Total Quotation Amount: ")
    r1.bold = True
    r2 = p.add_run(PH("Amount"))
    r2.bold = True
    r2.italic = True
    from doc_helpers import PLACEHOLDER_COLOR
    r2.font.color.rgb = PLACEHOLDER_COLOR

    h2(doc, "Commercial Terms")
    field(doc, "Payment Terms", "e.g. 50% upfront, 30% at midpoint, 20% on handover")
    field(doc, "Delivery Milestones", PH("Milestone Plan"), True)
    field(doc, "Assumptions", PH("Project assumptions"), True)
    field(doc, "Exclusions", PH("Out-of-scope items"), True)

    h2(doc, "Acceptance")
    signature_block(doc, "Client", PH("Client Name"), f"Authorized by {COMPANY['name']}", COMPANY["founder"])

    save(doc, "templates/business-documents/QUOTATION_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 4. INVOICE
# ─────────────────────────────────────────────────────────────────
def invoice():
    doc = new_document()
    doc_title(doc, "Invoice")
    placeholder_legend(doc)

    h2(doc, "Invoice Details")
    field(doc, "Invoice Number", PH("Invoice Number"), True)
    field(doc, "Invoice Date", PH("Invoice Date"), True)
    field(doc, "Due Date", PH("Due Date"), True)
    field(doc, "Currency", PH("KES / USD / Other"), True)
    field(doc, "Purchase Order (if any)", PH("PO Number"), True)

    h2(doc, "From")
    field(doc, "Company", COMPANY["name"])
    field(doc, "Registered Address", PH("Registered Address"), True)
    field(doc, "Email", COMPANY["email"])
    field(doc, "Phone", COMPANY["phone"])
    field(doc, "Company Registration Number", PH("Company Registration Number"), True)
    field(doc, "KRA PIN", PH("KRA PIN"), True)

    h2(doc, "Bill To")
    field(doc, "Client Name", PH("Client Name"), True)
    field(doc, "Company", PH("Client Company"), True)
    field(doc, "Billing Address", PH("Client Billing Address"), True)
    field(doc, "Email", PH("Client Billing Email"), True)

    h2(doc, "Invoice Items")
    make_table(
        doc,
        ["Item", "Description", "Qty", "Unit Price", "Amount"],
        [
            ["1", "[Milestone / Service]", "[Qty]", "[Amount]", "[Amount]"],
            ["2", "[Milestone / Service]", "[Qty]", "[Amount]", "[Amount]"],
            ["3", "[Milestone / Service]", "[Qty]", "[Amount]", "[Amount]"],
        ],
        align_right_cols={2, 3, 4},
    )
    field(doc, "Subtotal", PH("Amount"), True)
    field(doc, "VAT / Tax", f"({PH('Rate')}%) {PH('Amount')}", True)
    field(doc, "Less Withholding Tax (if applicable)", PH("Amount"), True)
    p = doc.add_paragraph()
    r1 = p.add_run("Total Due: ")
    r1.bold = True
    r2 = p.add_run(PH("Amount"))
    r2.bold = True
    r2.italic = True
    from doc_helpers import PLACEHOLDER_COLOR
    r2.font.color.rgb = PLACEHOLDER_COLOR

    h2(doc, "Payment Instructions")
    field(doc, "Account Name", PH("Account Name"), True)
    field(doc, "Bank Name", PH("Bank Name"), True)
    field(doc, "Account Number", PH("Bank Account Number"), True)
    field(doc, "Branch", PH("Branch"), True)
    field(doc, "SWIFT Code", PH("SWIFT Code"), True)
    field(doc, "Mobile Money Option (if applicable)", PH("Paybill / Till / Number"), True)

    h2(doc, "Notes")
    field(doc, "Late payment terms", PH("Late payment terms"), True)
    field(doc, "Support contact", COMPANY["email"])

    save(doc, "templates/business-documents/INVOICE_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 5. RECEIPT
# ─────────────────────────────────────────────────────────────────
def receipt():
    doc = new_document()
    doc_title(doc, "Payment Receipt")
    placeholder_legend(doc)

    h2(doc, "Receipt Details")
    field(doc, "Receipt Number", PH("Receipt Number"), True)
    field(doc, "Date", PH("Receipt Date"), True)
    field(doc, "Currency", PH("KES / USD / Other"), True)

    h2(doc, "Received By")
    body(doc, COMPANY["name"], bold=True)

    h2(doc, "Received From")
    field(doc, "Client Name", PH("Client Name"), True)
    field(doc, "Company", PH("Client Company"), True)
    field(doc, "Contact", PH("Client Contact"), True)

    h2(doc, "Payment Information")
    field(doc, "Amount Received", PH("Amount"), True)
    field(doc, "Payment Method", PH("Bank Transfer / Mobile Money / Cash / Card"), True)
    field(doc, "Reference Number", PH("Transaction Reference"), True)
    field(doc, "Related Invoice Number", PH("Invoice Number"), True)
    field(doc, "Payment Date", PH("Payment Date"), True)

    h2(doc, "Description")
    body(doc, PH("Description of payment purpose"), italic=True)

    h2(doc, "Confirmation")
    signature_block(doc, "Received in Full By", COMPANY["founder"])
    field(doc, "For queries", COMPANY["email"])

    save(doc, "templates/business-documents/RECEIPT_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 6. PROPOSAL
# ─────────────────────────────────────────────────────────────────
def proposal():
    doc = new_document()
    doc_title(doc, PH("Project Name") + " Proposal", f"Prepared for {PH('Client Name / Organization')}")
    placeholder_legend(doc)

    field(doc, "Prepared By", COMPANY["name"])
    field(doc, "Date", PH("Date"), True)
    field(doc, "Version", PH("Version Number"), True)

    h2(doc, "1. Executive Summary")
    body(doc, PH("Summarize the client's challenge, the strategic opportunity, and the "
                 "proposed solution in 3-5 sentences."), italic=True)

    h2(doc, "2. Client Context")
    field(doc, "Client Industry", PH("Industry"), True)
    field(doc, "Current Situation", PH("Current state / problem"), True)
    field(doc, "Business Objectives", PH("Primary objectives"), True)

    h2(doc, "3. Proposed Solution")
    h3(doc, "3.1 Scope of Work")
    body(doc, PH("Describe in-scope workstreams and deliverables."), italic=True)
    h3(doc, "3.2 Service Model")
    field(doc, "Primary Service", PH("Custom Software / Consulting / Transformation / Cloud / QA / Support"), True)
    field(doc, "Engagement Model", PH("Fixed-Price / Agile Retainer / Dedicated Team"), True)
    h3(doc, "3.3 Delivery Approach")
    for step in ["Discovery and planning", "Design and architecture", "Build and validation",
                 "Deployment and handover", "Post-launch support"]:
        numbered(doc, step)

    h2(doc, "4. Deliverables")
    make_table(
        doc,
        ["Deliverable", "Description", "Owner", "Target Date"],
        [
            ["[Deliverable 1]", "[Description]", "[Owner]", "[Date]"],
            ["[Deliverable 2]", "[Description]", "[Owner]", "[Date]"],
            ["[Deliverable 3]", "[Description]", "[Owner]", "[Date]"],
        ],
    )

    h2(doc, "5. Timeline and Milestones")
    field(doc, "Estimated Duration", PH("X weeks / months"), True)
    field(doc, "Key Milestones", PH("Milestone list"), True)
    field(doc, "Dependencies", PH("Dependencies"), True)

    h2(doc, "6. Investment")
    make_table(
        doc,
        ["Cost Component", "Amount"],
        [
            ["[Component 1]", "[Amount]"],
            ["[Component 2]", "[Amount]"],
            ["[Component 3]", "[Amount]"],
        ],
        align_right_cols={1},
    )
    field(doc, "Payment Terms", PH("Payment terms"), True)

    h2(doc, "7. Assumptions and Exclusions")
    field(doc, "Assumptions", PH("Assumptions"), True)
    field(doc, "Exclusions", PH("Out-of-scope items"), True)

    h2(doc, "8. Team and Governance")
    field(doc, "Engagement Lead", COMPANY["founder"])
    field(doc, "Technical Lead", COMPANY["founder"])
    field(doc, "Reporting Cadence", PH("Weekly / Bi-weekly"), True)
    field(doc, "Communication Channels", "Email / WhatsApp / Video Call")

    h2(doc, "9. Acceptance")
    body(doc, "If this proposal meets your expectations, please confirm approval below.")
    signature_block(doc, "Client", PH("Client Name"), f"Authorized by {COMPANY['name']}", COMPANY["founder"])

    save(doc, "templates/business-documents/PROPOSAL_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 7. NDA
# ─────────────────────────────────────────────────────────────────
def nda():
    doc = new_document()
    doc_title(doc, "Mutual Non-Disclosure Agreement")
    legal_notice(doc)
    placeholder_legend(doc)

    h2(doc, "1. Parties")
    body(doc, "This Mutual Non-Disclosure Agreement (“NDA”) is made between:")
    bullet(doc, f"Disclosing/Receiving Party A: {COMPANY['name']}")
    bullet(doc, f"Disclosing/Receiving Party B: {PH('Client Legal Name')}")
    field(doc, "Effective Date", PH("Effective Date"), True)

    h2(doc, "2. Purpose")
    body(doc, f"The parties wish to explore a potential business relationship concerning "
              f"{PH('Project/Opportunity Name')} and may share confidential information.")

    h2(doc, "3. Confidential Information")
    body(doc, "Confidential information includes non-public business, technical, financial, "
              "product, customer, and operational information disclosed in any form.")

    h2(doc, "4. Obligations")
    body(doc, "Each party agrees to:")
    for b in ["Use confidential information only for the stated purpose",
              "Protect confidential information with reasonable care",
              "Limit access to authorized personnel on a need-to-know basis",
              "Not disclose confidential information to third parties without written consent"]:
        bullet(doc, b)

    h2(doc, "5. Exclusions")
    body(doc, "Confidential information does not include information that:")
    for b in ["Is publicly available without breach",
              "Was already known before disclosure",
              "Is independently developed without reference to disclosed data",
              "Is required to be disclosed by law (with notice where permitted)"]:
        bullet(doc, b)

    h2(doc, "6. Term and Survival")
    field(doc, "NDA Term", PH("X years"), True)
    field(doc, "Confidentiality survives termination/expiry for", PH("X years"), True)

    h2(doc, "7. Return / Destruction")
    body(doc, "Upon written request, the receiving party will return or securely destroy "
              "confidential information, subject to legal retention obligations.")

    h2(doc, "8. No License")
    body(doc, "No intellectual property rights or licenses are granted except the limited "
              "right to evaluate the opportunity.")

    h2(doc, "9. Remedies")
    body(doc, "Unauthorized disclosure may cause irreparable harm; parties may seek injunctive "
              "relief and other legal remedies.")

    h2(doc, "10. Governing Law")
    body(doc, f"This NDA is governed by the laws of {PH('Governing Law/Jurisdiction')}.")

    h2(doc, "11. Signatures")
    signature_block(doc, f"Party A ({COMPANY['name']})", COMPANY["founder"],
                     f"Party B ({PH('Client Legal Name')})", PH("Name, Title"))

    save(doc, "templates/legal/NDA_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 8. SERVICE AGREEMENT
# ─────────────────────────────────────────────────────────────────
def service_agreement():
    doc = new_document()
    doc_title(doc, "Service Agreement")
    legal_notice(doc)
    placeholder_legend(doc)

    h2(doc, "1. Parties")
    body(doc, "This Service Agreement is made between:")
    bullet(doc, f"Service Provider: {COMPANY['name']} (“Provider”)")
    bullet(doc, f"Client: {PH('Client Legal Name')} (“Client”)")
    field(doc, "Effective Date", PH("Effective Date"), True)

    h2(doc, "2. Services")
    body(doc, "Provider agrees to perform the services described in Schedule A (Scope of "
              "Work), including deliverables, milestones, and timelines.")

    h2(doc, "3. Term")
    field(doc, "Start Date", PH("Start Date"), True)
    field(doc, "End Date", PH("End Date or Ongoing"), True)
    field(doc, "Renewal Terms", PH("Renewal Terms"), True)

    h2(doc, "4. Fees and Payment")
    field(doc, "Total Fees", PH("Amount/Currency"), True)
    field(doc, "Payment Schedule", PH("Milestone/Monthly terms"), True)
    field(doc, "Invoicing", PH("Invoice frequency"), True)
    field(doc, "Payment Due", PH("Net X days"), True)
    body(doc, "Taxes: Client is responsible for applicable taxes unless otherwise stated.")

    h2(doc, "5. Change Requests")
    body(doc, "Any scope changes must be documented and approved in writing by both parties, "
              "including cost and timeline impact.")

    h2(doc, "6. Client Responsibilities")
    body(doc, "Client shall:")
    for b in ["Provide required content, approvals, and access on time",
              "Assign a decision-making contact person",
              "Review deliverables within agreed timelines"]:
        bullet(doc, b)

    h2(doc, "7. Intellectual Property")
    for b in ["Pre-existing IP remains with its original owner",
              "Upon full payment, project-specific deliverables transfer to Client unless otherwise stated",
              "Provider may reuse generic know-how, frameworks, and non-confidential components"]:
        bullet(doc, b)

    h2(doc, "8. Confidentiality")
    body(doc, "Both parties must keep confidential information secure and use it only for "
              "this engagement.")

    h2(doc, "9. Warranties and Support")
    field(doc, "Initial warranty period", PH("e.g. 30 days"), True)
    body(doc, "Ongoing support is governed by separate maintenance/support terms if purchased.")

    h2(doc, "10. Limitation of Liability")
    body(doc, PH("Insert legal-approved liability limitation and exclusion clause."), italic=True)

    h2(doc, "11. Termination")
    field(doc, "Notice period", PH("X days"), True)
    field(doc, "Payment obligations on termination", PH("Terms"), True)
    field(doc, "Transition support", PH("Terms"), True)

    h2(doc, "12. Governing Law")
    body(doc, f"This agreement is governed by the laws of {PH('Governing Law/Jurisdiction')}.")

    h2(doc, "13. Signatures")
    signature_block(doc, "Provider Representative", COMPANY["founder"],
                     "Client Representative", PH("Name, Title"))

    h2(doc, "Schedule A: Scope of Work")
    body(doc, PH("Detailed scope, deliverables, timeline, and assumptions."), italic=True)

    save(doc, "templates/legal/SERVICE_AGREEMENT_TEMPLATE.docx")


# ─────────────────────────────────────────────────────────────────
# 9. TERMS AND CONDITIONS
# ─────────────────────────────────────────────────────────────────
def terms_and_conditions():
    doc = new_document()
    doc_title(doc, "Terms and Conditions", "General business terms for engagements with Nakola Expert Systems")
    legal_notice(doc)
    placeholder_legend(doc)

    h2(doc, "1. Acceptance of Terms")
    body(doc, f"By engaging {COMPANY['name']} or using {COMPANY['website']}, you agree to "
              "these Terms and Conditions.")

    h2(doc, "2. Services")
    body(doc, f"{COMPANY['name']} provides software, consulting, digital transformation, "
              "cloud, QA, and support services under separate signed project agreements.")

    h2(doc, "3. Use of Website")
    body(doc, "Users agree not to:")
    for b in ["Violate applicable laws or regulations", "Attempt unauthorized access to our systems",
              "Misuse website content or systems"]:
        bullet(doc, b)

    h2(doc, "4. Intellectual Property")
    body(doc, f"All website content, branding, and proprietary materials are owned by "
              f"{COMPANY['name']} or licensed to us unless otherwise stated.")

    h2(doc, "5. Quotes and Proposals")
    field(doc, "Quotation validity", PH("X days"), True)
    body(doc, "Final scope, pricing, and timelines are governed by signed agreements.")

    h2(doc, "6. Payments")
    body(doc, "Payment terms, late fees, and taxes are governed by the signed commercial "
              "document for each engagement (quotation, invoice, or service agreement).")

    h2(doc, "7. Disclaimer")
    body(doc, "Website content is for general informational purposes and may change without notice.")

    h2(doc, "8. Limitation of Liability")
    body(doc, PH("Insert legal-approved liability limitation and exclusion clause."), italic=True)

    h2(doc, "9. Third-Party Links")
    body(doc, "We are not responsible for external third-party websites linked from our website.")

    h2(doc, "10. Termination / Suspension")
    body(doc, "We may suspend access where misuse, breach, or legal risk is identified.")

    h2(doc, "11. Governing Law")
    body(doc, f"These terms are governed by the laws of {PH('Governing Law/Jurisdiction')}.")

    h2(doc, "12. Contact")
    field(doc, "Email", COMPANY["email"])
    field(doc, "Address", PH("Registered Address"), True)

    save(doc, "templates/legal/TERMS_AND_CONDITIONS_TEMPLATE.docx")


if __name__ == "__main__":
    company_profile()
    letterhead()
    quotation()
    invoice()
    receipt()
    proposal()
    nda()
    service_agreement()
    terms_and_conditions()
    print("\nAll 9 documents generated successfully.")
